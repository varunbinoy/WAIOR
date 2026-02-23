
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# ===============================
# CONFIGURATION
# ===============================

PROJECT_ROOT = Path(".")
SRC_DIR = PROJECT_ROOT / "src"
OUTPUT_PATH = PROJECT_ROOT / "outputs" / "WAI_Code_Appendix.docx"

# Update filenames here if needed
MODULES = [
    {
        "file": "01_load_data.py",
        "title": "1. Data Ingestion Module",
        "description": (
            "Reads the input Excel dataset, standardizes identifiers, removes duplicates, "
            "and exports normalized CSVs (courses, enrollments, etc.) for downstream optimization."
        )
    },
    {
        "file": "02c_sectioning_ilp.py",
        "title": "2. Sectioning Module (ILP-Based)",
        "description": (
            "Implements CP-SAT based section splitting for oversubscribed courses. "
            "Enforces 25–70 capacity bounds and minimizes deviation from a global A/B bucket assignment."
        )
    },
    {
        "file": "03_build_slots.py",
        "title": "3. Slot Construction Module",
        "description": (
            "Constructs the 10-week timetable structure (400 logical slots) and embeds "
            "dynamic classroom availability (Weeks 1–4: 10 rooms; Weeks 5–10: 4 rooms)."
        )
    },
    {
        "file": "04_solve_term_schedule.py",
        "title": "4. Core Term Solver",
        "description": (
            "Builds the CP-SAT optimization model with decision variables x(section, slot). "
            "Enforces room capacity, student conflicts, faculty conflicts, session bounds (18–20), "
            "and faculty week-split logic. Maximizes total scheduled sessions."
        )
    },
    {
        "file": "61_solve_contingent_days_min.py",
        "title": "5. Contingent Solver Module",
        "description": (
            "Schedules remaining deficit sessions using contingent days (7 slots/day, 10 rooms). "
            "Minimizes number of contingent days activated while maintaining all conflict constraints."
        )
    },
]

# ===============================
# STYLING FUNCTIONS
# ===============================

def set_document_style(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")

def add_code_block(doc, code_text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")

# ===============================
# MAIN LOGIC
# ===============================

def main():
    doc = Document()
    set_document_style(doc)

    doc.add_heading("WAI Timetable Optimization — Code Appendix", level=1)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(
        "This appendix contains the primary Python modules used in the timetable optimization pipeline."
    )

    for module in MODULES:
        doc.add_page_break()
        doc.add_heading(module["title"], level=2)

        file_path = SRC_DIR / module["file"]

        doc.add_paragraph("Purpose:")
        doc.add_paragraph(module["description"])

        if not file_path.exists():
            doc.add_paragraph(f"⚠ ERROR: File not found → {file_path}")
            continue

        code_text = file_path.read_text(encoding="utf-8", errors="ignore")

        doc.add_paragraph("Code:")
        add_code_block(doc, code_text)

    OUTPUT_PATH.parent.mkdir(exist_ok=True)
    doc.save(OUTPUT_PATH)

    print(f"\n✅ Code appendix generated at:\n{OUTPUT_PATH.resolve()}")

if __name__ == "__main__":
    main()
